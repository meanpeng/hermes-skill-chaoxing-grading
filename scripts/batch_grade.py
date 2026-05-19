#!/usr/bin/env python3
"""
Prepare Chaoxing/Xuexitong assignment attachments for agent-assisted grading.

The script extracts report text from .docx/.doc files, expands common nested
"one zip per student" exports, and writes a CSV material list. It does not
assign grades by itself.
"""

import argparse
import csv
import os
import struct
import sys
import zipfile


REPORT_TEMPLATE_KEYWORDS = ("V2024", "模板", "template")
STUDENT_ARCHIVE_DIR = "__unzipped__"
DEFAULT_MIN_CHARS = 50

SECTION_KEYWORDS = [
    ("basic", ["课程名称", "项目名称", "实验/实训报告", "学生姓名", "学号"]),
    ("purpose_or_type", ["实验目的", "实验类型", "设计性实验", "验证性实验", "项目名称"]),
    ("environment", ["实验环境", "软件环境", "硬件设备", "依赖库", "Python", "numpy", "matplotlib"]),
    ("process", ["实验步骤", "内容及过程记录", "关键函数", "代码解释", "算法流程", "训练过程"]),
    ("result", ["实验结果", "结果与分析", "运行结果", "可视化", "截图", "输出结果"]),
    ("summary_or_key_points", ["实验总结", "总结", "心得", "体会", "分析", "关键点", "任务要求", "作业要求", "收获"]),
]


def open_zip_with_encoding(zip_path, encoding="gbk"):
    """Open a zip with GBK fallback while honoring UTF-8 filename flags."""
    with zipfile.ZipFile(zip_path) as probe:
        has_utf8_names = any(info.flag_bits & 0x800 for info in probe.infolist())
    if has_utf8_names:
        return zipfile.ZipFile(zip_path)
    try:
        return zipfile.ZipFile(zip_path, metadata_encoding=encoding)
    except TypeError:
        return zipfile.ZipFile(zip_path)


def safe_extract_zip(zip_path, output_dir, encoding="gbk"):
    """Extract a zip safely, preserving UTF-8 or GBK names where possible."""
    base = os.path.abspath(output_dir)
    os.makedirs(base, exist_ok=True)
    with open_zip_with_encoding(zip_path, encoding) as zf:
        for member in zf.infolist():
            target = os.path.abspath(os.path.join(base, member.filename))
            if os.path.commonpath([base, target]) != base:
                raise ValueError(f"Unsafe zip path: {member.filename}")
            zf.extract(member, base)


def prepare_student_dir(student_dir):
    """Return a directory to inspect, expanding nested student zips if needed."""
    if any(name.lower().endswith((".docx", ".doc")) for name in os.listdir(student_dir)):
        return student_dir, False

    zip_files = [
        name for name in os.listdir(student_dir)
        if name.lower().endswith(".zip") and os.path.isfile(os.path.join(student_dir, name))
    ]
    if not zip_files:
        return student_dir, False

    expanded = os.path.join(student_dir, STUDENT_ARCHIVE_DIR)
    os.makedirs(expanded, exist_ok=True)
    for zip_name in zip_files:
        safe_extract_zip(os.path.join(student_dir, zip_name), expanded)
    return expanded, True


def docx_text(path):
    from docx import Document

    doc = Document(path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return "\n".join(texts)


def doc_text(path):
    import olefile

    ole = olefile.OleFileIO(path)
    if not ole.exists("WordDocument"):
        return ""
    raw = ole.openstream("WordDocument").read()
    text_parts, current = [], []
    i = 0
    while i < len(raw) - 1:
        char = struct.unpack("<H", raw[i:i + 2])[0]
        if (
            0x20 <= char <= 0x7E
            or 0x4E00 <= char <= 0x9FFF
            or 0x3000 <= char <= 0x303F
            or 0xFF00 <= char <= 0xFFEF
            or char in (0x0D, 0x0A, 0x09)
        ):
            current.append(chr(char))
        else:
            if len(current) > 3:
                text_parts.append("".join(current))
            current = []
        i += 2
    if len(current) > 3:
        text_parts.append("".join(current))
    return "\n".join(text_parts)


def candidate_reports(student_dir):
    files = os.listdir(student_dir)
    docxs = [
        name for name in files
        if name.lower().endswith(".docx")
        and not any(keyword.lower() in name.lower() for keyword in REPORT_TEMPLATE_KEYWORDS)
    ]
    if not docxs:
        docxs = [name for name in files if name.lower().endswith(".docx")]
    docs = [name for name in files if name.lower().endswith(".doc") and not name.lower().endswith(".docx")]
    return [(os.path.join(student_dir, name), "docx") for name in docxs] + [
        (os.path.join(student_dir, name), "doc") for name in docs
    ]


def get_report_text(student_dir, min_chars=DEFAULT_MIN_CHARS):
    """
    Return (text, path, format, status).

    status values:
    - ok: readable report with enough text for normal review
    - too_short: readable report found, but text is below min_chars
    - unreadable: report-like files exist but could not be parsed
    - missing_report: no .doc/.docx report found
    """
    reports = candidate_reports(student_dir)
    if not reports:
        return "", None, "none", "missing_report"

    best_short = None
    unreadable = False
    for path, fmt in reports:
        try:
            text = docx_text(path) if fmt == "docx" else doc_text(path)
        except Exception:
            unreadable = True
            continue

        if len(text) >= min_chars:
            return text, path, fmt, "ok"
        if text and (best_short is None or len(text) > len(best_short[0])):
            best_short = (text, path, fmt)

    if best_short:
        text, path, fmt = best_short
        return text, path, fmt, "too_short"
    return "", None, "none", "unreadable" if unreadable else "missing_report"


def count_report_images(format_type, report_path):
    """Count report images. .doc image count is estimated from file size."""
    if not report_path:
        return 0
    if format_type == "docx":
        try:
            from docx import Document
            doc = Document(report_path)
            return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
        except Exception:
            return 0
    if format_type == "doc":
        size = os.path.getsize(report_path)
        if size > 500000:
            return max(3, min(15, (size - 200000) // 150000))
    return 0


def preview_text(full_text, limit=220):
    """Return a compact text preview for orientation only; not for grading."""
    collapsed = " ".join(line.strip() for line in full_text.splitlines() if line.strip())
    return collapsed[:limit]


def optional_metrics(full_text):
    """Return rough signals for optional metric-assisted screening."""
    section_count = sum(
        1 for _, keywords in SECTION_KEYWORDS
        if any(keyword in full_text for keyword in keywords)
    )
    has_key_requirement = any(keyword in full_text for keyword in [
        "任务要求", "作业要求", "实验要求", "关键点", "关键步骤",
        "实现", "代码", "运行", "结果", "截图", "输出", "分析",
        "总结", "心得", "体会", "收获",
    ])
    return min(6, section_count), has_key_requirement


def analyze_student(extract_dir, student_dir_name, include_metrics=False, min_chars=DEFAULT_MIN_CHARS):
    student_dir = os.path.join(extract_dir, student_dir_name)
    if not os.path.isdir(student_dir):
        return None
    inspect_dir, expanded_zip = prepare_student_dir(student_dir)

    parts = student_dir_name.split("-", 1)
    student_id = parts[0] if parts else student_dir_name
    student_name = parts[1] if len(parts) > 1 else "unknown"

    full_text, report_path, fmt, status = get_report_text(inspect_dir, min_chars=min_chars)
    img_count = count_report_images(fmt, report_path)

    result = {
        "student_id": student_id,
        "student_name": student_name,
        "dir_name": student_dir_name,
        "status": status,
        "format": fmt,
        "char_count": len(full_text),
        "img_count": img_count,
        "expanded_zip": expanded_zip,
        "report_path": report_path or "",
        "preview": preview_text(full_text),
    }
    if include_metrics:
        section_count, has_key_requirement = optional_metrics(full_text)
        result.update({
            "section_signal": section_count,
            "key_requirement_signal": has_key_requirement,
        })
    return result


def resolve_extract_dir(base_dir):
    """Support common export layouts: base/docx_files, base/extracted, or base."""
    for name in ("docx_files", "extracted"):
        candidate = os.path.join(base_dir, name)
        if os.path.isdir(candidate):
            return candidate
    return base_dir


def ensure_student_dirs(extract_dir):
    """If Chaoxing exported one zip per student at top level, expand them."""
    for name in os.listdir(extract_dir):
        path = os.path.join(extract_dir, name)
        if not os.path.isfile(path) or not name.lower().endswith(".zip"):
            continue
        student_name = os.path.splitext(name)[0]
        target = os.path.join(extract_dir, student_name)
        os.makedirs(target, exist_ok=True)
        safe_extract_zip(path, target)


def sample_results(results, sample_size):
    if not sample_size or sample_size >= len(results):
        return results
    if sample_size <= 0:
        return []
    if sample_size == 1:
        return [results[0]]

    last = len(results) - 1
    indexes = sorted({round(i * last / (sample_size - 1)) for i in range(sample_size)})
    return [results[i] for i in indexes]


def parse_args():
    parser = argparse.ArgumentParser(
        description="Prepare Chaoxing assignment materials for agent-assisted grading."
    )
    parser.add_argument(
        "--base-dir",
        default=os.getcwd(),
        help="Extracted assignment directory, or a parent containing docx_files/extracted.",
    )
    parser.add_argument(
        "--output-csv",
        default=None,
        help="CSV output path. Defaults to base-dir/grading_materials.csv or grading_metrics.csv.",
    )
    parser.add_argument(
        "--mode",
        choices=("materials", "metrics"),
        default="materials",
        help="materials only, or metrics with rough section/key-requirement signals.",
    )
    parser.add_argument(
        "--sample",
        type=int,
        default=0,
        help="Print only a representative sample in the console; CSV still includes all rows.",
    )
    parser.add_argument(
        "--min-chars",
        type=int,
        default=DEFAULT_MIN_CHARS,
        help=f"Minimum extracted text length for status=ok. Default: {DEFAULT_MIN_CHARS}.",
    )
    return parser.parse_args()


def print_rows(rows, include_metrics):
    if include_metrics:
        print("MODE: metrics (rough signals only; still read submissions before grading)")
        print(f"{'student_id':<12} {'name':<12} {'status':<14} {'fmt':<5} {'chars':>6} {'imgs':>4} {'sec':>5} {'key':>5}  report")
    else:
        print("MODE: materials (organize files only; no automatic grading)")
        print(f"{'student_id':<12} {'name':<12} {'status':<14} {'fmt':<5} {'chars':>6} {'imgs':>4}  report")
    print("-" * 96)
    for row in rows:
        if include_metrics:
            key_requirement = "yes" if row["key_requirement_signal"] else "no"
            print(
                f"{row['student_id']:<12} {row['student_name']:<12} {row['status']:<14} "
                f"{row['format']:<5} {row['char_count']:>6} {row['img_count']:>4} "
                f"{row['section_signal']:>4}/6 {key_requirement:>5}  {row['report_path']}"
            )
        else:
            print(
                f"{row['student_id']:<12} {row['student_name']:<12} {row['status']:<14} "
                f"{row['format']:<5} {row['char_count']:>6} {row['img_count']:>4}  {row['report_path']}"
            )


def main():
    args = parse_args()
    base_dir = os.path.abspath(args.base_dir)
    include_metrics = args.mode == "metrics"
    default_name = "grading_metrics.csv" if include_metrics else "grading_materials.csv"
    output_csv = os.path.abspath(args.output_csv or os.path.join(base_dir, default_name))

    extract_dir = resolve_extract_dir(base_dir)
    if not os.path.isdir(extract_dir):
        print(f"ERROR: {extract_dir} not found. Unzip first.", file=sys.stderr)
        return 1

    ensure_student_dirs(extract_dir)

    results = []
    for name in sorted(os.listdir(extract_dir)):
        row = analyze_student(extract_dir, name, include_metrics=include_metrics, min_chars=args.min_chars)
        if row:
            results.append(row)

    display_rows = sample_results(sorted(results, key=lambda row: row["student_id"]), args.sample)
    print_rows(display_rows, include_metrics=include_metrics)
    if args.sample:
        print(f"\nShowing {len(display_rows)} of {len(results)} submissions. CSV includes all rows.")
    else:
        print(f"\nTotal submissions: {len(results)}")

    status_counts = {}
    for row in results:
        status_counts[row["status"]] = status_counts.get(row["status"], 0) + 1
    print("Status counts: " + ", ".join(f"{k}={v}" for k, v in sorted(status_counts.items())))

    with open(output_csv, "w", newline="", encoding="utf-8-sig") as file:
        fieldnames = [
            "student_id", "student_name", "dir_name", "status", "format",
            "char_count", "img_count", "expanded_zip", "report_path", "preview",
        ]
        if include_metrics:
            fieldnames.extend(["section_signal", "key_requirement_signal"])
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(results)
    print(f"CSV saved: {output_csv}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
